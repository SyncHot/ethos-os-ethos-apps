"""
EthOS -- Document Anonymizer
Anonymize medical PDF/DOCX documents using local Bielik LLM.
Extracts text, identifies PII via LLM, replaces with placeholders.

Endpoints:
  POST /api/doc-anonymizer/upload           -- upload and anonymize a file
  GET  /api/doc-anonymizer/jobs             -- list anonymization jobs
  GET  /api/doc-anonymizer/download/<job_id> -- download anonymized file
  GET  /api/doc-anonymizer/preview/<job_id>/<which> -- inline preview (original|anonymized)
  DELETE /api/doc-anonymizer/job/<job_id>   -- delete a job
  POST /api/doc-anonymizer/jobs/delete-batch -- batch-delete multiple jobs
  GET  /api/doc-anonymizer/status           -- app + AI dependency status
  POST /api/doc-anonymizer/install          -- install (via register_pkg_routes)
  POST /api/doc-anonymizer/uninstall        -- uninstall
  GET  /api/doc-anonymizer/pkg-status       -- package status

SocketIO events:
  anon_progress  -- real-time anonymization progress updates
"""

import json
import os
import re
import sys
import time
import uuid
import logging
import threading
import unicodedata

from flask import Blueprint, request, jsonify, send_file, g

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from utils import register_pkg_routes, get_username_or
from host import data_path, host_run

log = logging.getLogger('ethos.doc_anonymizer')

doc_anonymizer_bp = Blueprint('doc_anonymizer', __name__,
                              url_prefix='/api/doc-anonymizer')

# -- Paths ------------------------------------------------------------------
_JOBS_DIR = data_path('doc_anonymizer_jobs')


def _ensure_jobs_dir():
    os.makedirs(_JOBS_DIR, exist_ok=True)


def _ensure_deps():
    """Install all required pip packages and system tools if missing."""
    missing_pip = []
    try:
        import fitz  # noqa: F401
    except ImportError:
        missing_pip.append('PyMuPDF')
    try:
        import PyPDF2  # noqa: F401
    except ImportError:
        missing_pip.append('PyPDF2')
    try:
        from PIL import Image  # noqa: F401
    except ImportError:
        missing_pip.append('Pillow')
    try:
        import pytesseract  # noqa: F401
    except ImportError:
        missing_pip.append('pytesseract')
    try:
        import docx  # noqa: F401
    except ImportError:
        missing_pip.append('python-docx')
    try:
        import reportlab  # noqa: F401
    except ImportError:
        missing_pip.append('reportlab')
    try:
        import spacy  # noqa: F401
    except ImportError:
        missing_pip.append('spacy')

    if missing_pip:
        pkgs = ' '.join(missing_pip)
        log.info('[doc_anonymizer] Installing pip deps: %s', pkgs)
        host_run(f'/opt/ethos/venv/bin/pip install --quiet {pkgs}', timeout=120)

    # Download spaCy Polish model if not available
    try:
        import spacy
        spacy.load('pl_core_news_lg')
    except Exception:
        log.info('[doc_anonymizer] Downloading spaCy Polish model (pl_core_news_lg)')
        host_run('/opt/ethos/venv/bin/python -m spacy download pl_core_news_lg', timeout=300)

    # Download NER delta weights if neither full model nor delta exist
    _ensure_ner_delta()

    import shutil
    if not shutil.which('pdftotext'):
        log.info('[doc_anonymizer] Installing poppler-utils')
        host_run('apt-get install -y -qq poppler-utils && apt-get clean', timeout=60)

    if not shutil.which('tesseract'):
        log.info('[doc_anonymizer] Installing Tesseract OCR + Polish language pack')
        host_run('apt-get install -y -qq tesseract-ocr tesseract-ocr-pol && apt-get clean', timeout=120)


_NER_DELTA_URL = (
    'https://raw.githubusercontent.com/SyncHot/ethos-os-ethos-apps/main/'
    'apps/doc-anonymizer/models/spacy_pii_ner_delta.tar.gz'
)


def _ensure_ner_delta():
    """Download NER delta model from GitHub release if not present locally."""
    models_dir = data_path('models')
    full_model = os.path.join(models_dir, 'spacy_pii_pl')
    delta_dir = os.path.join(models_dir, 'spacy_pii_ner_delta')

    if (os.path.isdir(full_model) and
            os.path.isfile(os.path.join(full_model, 'meta.json'))):
        return  # Full model already trained
    if (os.path.isdir(delta_dir) and
            os.path.isfile(os.path.join(delta_dir, 'meta.json'))):
        return  # Delta already present

    log.info('[doc_anonymizer] Downloading NER delta model...')
    os.makedirs(models_dir, exist_ok=True)
    tmp_tar = os.path.join(models_dir, '_ner_delta.tar.gz')
    try:
        import urllib.request
        urllib.request.urlretrieve(_NER_DELTA_URL, tmp_tar)
        import tarfile
        with tarfile.open(tmp_tar, 'r:gz') as tf:
            tf.extractall(models_dir)
        log.info('[doc_anonymizer] NER delta model installed to %s', delta_dir)
    except Exception as e:
        log.warning('[doc_anonymizer] Failed to download NER delta: %s', e)
    finally:
        if os.path.exists(tmp_tar):
            os.remove(tmp_tar)


# -- spaCy NER model (lazy-loaded singleton) --------------------------------

_spacy_nlp = None
_spacy_load_attempted = False


def _load_spacy_model():
    """Load spaCy PII model. Priority: full fine-tuned > delta NER + base > base only."""
    global _spacy_nlp, _spacy_load_attempted
    if _spacy_nlp is not None:
        return _spacy_nlp
    if _spacy_load_attempted:
        return None
    _spacy_load_attempted = True

    import spacy

    # 1) Full fine-tuned model (from training or first-run rebuild)
    pii_model_path = os.path.join(data_path('models'), 'spacy_pii_pl')
    if os.path.isdir(pii_model_path) and os.path.isfile(os.path.join(pii_model_path, 'meta.json')):
        try:
            _spacy_nlp = spacy.load(pii_model_path)
            log.info('[doc_anonymizer] Loaded fine-tuned spaCy PII model from %s', pii_model_path)
            return _spacy_nlp
        except Exception as e:
            log.warning('[doc_anonymizer] Failed to load fine-tuned model: %s', e)

    # 2) Delta NER weights + base model (shipped with app package, ~17MB)
    delta_path = os.path.join(data_path('models'), 'spacy_pii_ner_delta')
    delta_ner_path = os.path.join(delta_path, 'ner')
    if os.path.isdir(delta_ner_path) and os.path.isfile(os.path.join(delta_path, 'meta.json')):
        try:
            nlp = spacy.load('pl_core_news_lg')
            # Symlink base vocab vectors into delta so from_disk works
            delta_vocab = os.path.join(delta_ner_path, 'vocab')
            base_vocab_dir = str(nlp.path / 'vocab')
            _symlinks = []
            for fname in ('vectors', 'key2row'):
                src = os.path.join(base_vocab_dir, fname)
                dst = os.path.join(delta_vocab, fname)
                if os.path.exists(src) and not os.path.exists(dst):
                    os.symlink(src, dst)
                    _symlinks.append(dst)
            try:
                ner = nlp.get_pipe('ner')
                ner.from_disk(delta_ner_path)
                _spacy_nlp = nlp
                log.info('[doc_anonymizer] Loaded base pl_core_news_lg + NER delta from %s', delta_path)
                return _spacy_nlp
            finally:
                for lnk in _symlinks:
                    if os.path.islink(lnk):
                        os.remove(lnk)
        except Exception as e:
            log.warning('[doc_anonymizer] Failed to load NER delta: %s', e)

    # 3) Base Polish model only (no fine-tuning)
    try:
        _spacy_nlp = spacy.load('pl_core_news_lg')
        log.info('[doc_anonymizer] Loaded base spaCy model pl_core_news_lg')
        return _spacy_nlp
    except Exception as e:
        log.warning('[doc_anonymizer] Failed to load spaCy model: %s', e)

    # 4) Attempt to retrain if training script available
    train_script = os.path.join(os.environ.get('ETHOS_ROOT', '/opt/ethos'),
                                'tools', 'train_spacy_pii.py')
    if os.path.isfile(train_script):
        try:
            log.info('[doc_anonymizer] Attempting to train spaCy PII model...')
            from host import host_run
            out = host_run(
                f'/opt/ethos/venv/bin/python3 {train_script} --iterations 20',
                timeout=300
            )
            if os.path.isdir(pii_model_path):
                _spacy_nlp = spacy.load(pii_model_path)
                log.info('[doc_anonymizer] Trained and loaded spaCy PII model')
                return _spacy_nlp
        except Exception as e:
            log.warning('[doc_anonymizer] Training failed: %s', e)

    return None


# Medical title prefixes for LEKARZ classification
_DOCTOR_TITLE_PREFIXES = (
    'dr n. med.', 'dr hab. n. med.', 'prof. dr hab. n. med.',
    'prof. dr hab.', 'prof.', 'dr hab.', 'dr', 'lek. med.', 'lek.',
)


def _call_spacy_ner(text):
    """Detect PII entities using spaCy NER. Returns list of entity dicts.

    Maps spaCy labels to anonymizer categories:
      persName → IMIE_NAZWISKO or LEKARZ (if preceded by medical title)
      orgName  → NAZWA_PLACOWKI
      geogName → ADRES
      placeName → ADRES
    """
    nlp = _load_spacy_model()
    if nlp is None:
        log.warning('[doc_anonymizer] spaCy model not available')
        return []

    # Process in chunks if text is very long (spaCy default max is 1M chars)
    max_chunk = 100000
    entities = []
    seen = set()

    chunks = [text[i:i + max_chunk] for i in range(0, len(text), max_chunk)]
    for chunk in chunks:
        doc = nlp(chunk)
        for ent in doc.ents:
            ent_text = ent.text.strip()
            if not ent_text or len(ent_text) < 2:
                continue

            # Map spaCy label to our category
            if ent.label_ == 'persName':
                # Check if entity starts with a medical title
                ent_lower = ent_text.lower()
                is_doctor = any(ent_lower.startswith(tp) for tp in _DOCTOR_TITLE_PREFIXES)
                category = 'LEKARZ' if is_doctor else 'IMIE_NAZWISKO'
            elif ent.label_ == 'orgName':
                category = 'NAZWA_PLACOWKI'
            elif ent.label_ in ('geogName', 'placeName'):
                category = 'ADRES'
            else:
                continue  # Skip date, time, etc. (handled by regex)

            key = ent_text.lower()
            if key not in seen:
                seen.add(key)
                entities.append({'text': ent_text, 'category': category})

    log.info('[doc_anonymizer] spaCy NER found %d entities', len(entities))
    return entities


def _job_dir(job_id):
    return os.path.join(_JOBS_DIR, job_id)


def _get_username():
    return get_username_or('admin')


# -- Text extraction --------------------------------------------------------

def _extract_text_pdf(filepath):
    """Extract text from a PDF using pdftotext (poppler), fallback to PyPDF2, then OCR."""
    import subprocess
    try:
        result = subprocess.run(
            ['pdftotext', '-layout', filepath, '-'],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            raw = result.stdout
            pages = raw.split('\x0c')
            pages = [p for p in pages if p.strip()]
            if pages:
                return [_cleanup_pdf_text(p) for p in pages]
    except Exception as e:
        log.warning('[doc_anonymizer] pdftotext failed, using PyPDF2: %s', e)

    import PyPDF2
    pages = []
    with open(filepath, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text = page.extract_text() or ''
            pages.append(_cleanup_pdf_text(text))

    # If no text extracted (scanned PDF), try OCR
    has_text = any(p.strip() for p in pages)
    if not has_text:
        log.info('[doc_anonymizer] No text extracted, attempting OCR for %s', filepath)
        ocr_pages = _ocr_pdf(filepath)
        if ocr_pages:
            return ocr_pages

    return pages


def _ocr_pdf(filepath):
    """Extract text from a scanned PDF using Tesseract OCR."""
    try:
        import fitz
        import pytesseract
        from PIL import Image
        import io
    except ImportError as e:
        log.warning('[doc_anonymizer] OCR dependencies not available: %s', e)
        return None

    pages = []
    try:
        doc = fitz.open(filepath)
        for page_num, page in enumerate(doc):
            # Render page at 300 DPI for good OCR quality
            mat = fitz.Matrix(300 / 72, 300 / 72)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes('png')
            img = Image.open(io.BytesIO(img_data))

            # OCR with Polish + English language support
            text = pytesseract.image_to_string(img, lang='pol+eng', config='--psm 6')
            text = _normalize_ocr_text(text)
            pages.append(_cleanup_pdf_text(text))
            log.debug('[doc_anonymizer] OCR page %d: %d chars', page_num + 1, len(text))

        doc.close()
    except Exception as e:
        log.error('[doc_anonymizer] OCR failed: %s', e)
        return None

    has_text = any(p.strip() for p in pages)
    if not has_text:
        return None

    log.info('[doc_anonymizer] OCR extracted %d pages', len(pages))
    return pages


# Words that should stay ALL CAPS (abbreviations, headers)
_OCR_KEEP_UPPER = frozenset({
    'PESEL', 'NIP', 'REGON', 'KRS', 'PWZ', 'NFZ', 'ZUS', 'PIT', 'VAT',
    'KARTA', 'INFORMACYJNA', 'MR', 'CT', 'EKG', 'USG', 'RTG', 'MRI',
    'DNA', 'RNA', 'HIV', 'HCV', 'HBS', 'CRP', 'HDL', 'LDL', 'TSH',
    'BMI', 'EWUS', 'NZOZ', 'SP', 'ZOZ', 'II', 'III', 'IV', 'VI',
})


def _normalize_ocr_text(text):
    """Normalize OCR text: convert ALL CAPS person names to Title Case.

    OCR often outputs names in ALL CAPS (e.g. "ALICJA KOWALSKA").
    Regex patterns expect Title Case, so we normalize words that look
    like names while preserving known abbreviations.
    """
    lines = text.split('\n')
    normalized = []
    for line in lines:
        words = line.split()
        new_words = []
        for word in words:
            # Skip non-alpha, short words, known abbreviations
            stripped = word.strip('.,;:!?()[]/-')
            if (len(stripped) >= 3
                    and stripped.isupper()
                    and stripped.isalpha()
                    and stripped not in _OCR_KEEP_UPPER):
                # Convert to title case, preserve surrounding punctuation
                new_words.append(word.replace(stripped, stripped.title()))
            else:
                new_words.append(word)
        normalized.append(' '.join(new_words))
    return '\n'.join(normalized)


def _cleanup_pdf_text(text):
    """Reassemble fragmented text from PDF extraction.

    Tables in PDFs often produce single-char lines, scattered digits, and
    broken words.  This tries to rejoin them for better PII detection.
    """
    lines = text.split('\n')
    cleaned = []
    digit_buf = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            # Flush digit buffer on blank line
            if digit_buf:
                cleaned.append(''.join(digit_buf))
                digit_buf = []
            continue

        # Collect scattered single digits (likely PESEL fragments)
        if re.match(r'^\d{1,2}$', stripped):
            digit_buf.append(stripped)
            continue

        if digit_buf:
            cleaned.append(''.join(digit_buf))
            digit_buf = []

        # Collapse excessive whitespace within a line
        stripped = re.sub(r'\s{3,}', '  ', stripped)
        cleaned.append(stripped)

    if digit_buf:
        cleaned.append(''.join(digit_buf))

    return '\n'.join(cleaned)


def _extract_text_docx(filepath):
    """Extract text from a DOCX file paragraph by paragraph."""
    import docx
    doc = docx.Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return paragraphs


# -- Regex-based PII detection (fast, reliable for structured data) ---------

_MONTHS_PL = ('stycznia|lutego|marca|kwietnia|maja|czerwca|lipca|sierpnia|'
              'wrzesnia|września|pazdziernika|października|listopada|grudnia')

_REGEX_PATTERNS = [
    # PESEL: exactly 11 digits, not part of a longer number
    (re.compile(r'(?<!\d)\d{11}(?!\d)'), 'PESEL'),
    # Polish bank account (IBAN): 2+26 digits with spaces
    (re.compile(r'(?<!\d)\d{2}\s?\d{4}\s?\d{4}\s?\d{4}\s?\d{4}\s?\d{4}\s?\d{4}(?!\d)'), 'NR_KONTA'),
    # Document ID: 3 uppercase letters + 6 digits (e.g. CBA 123456)
    (re.compile(r'\b[A-Z]{3}\s?\d{6}\b'), 'NR_DOKUMENTU'),
    # Phone: Polish formats with +48
    (re.compile(r'(?:\+48|0048)[\s-]?\d{3}[\s-]?\d{3}[\s-]?\d{3}'), 'TELEFON'),
    # Phone: 9 digits with separators
    (re.compile(r'(?<!\d)\d{3}[\s-]\d{3}[\s-]\d{3}(?!\d)'), 'TELEFON'),
    # Phone: landline (2-digit area + 7 digits, e.g. "22 620 00 00")
    (re.compile(r'(?<!\d)\d{2}\s\d{3}\s\d{2}\s\d{2}(?!\d)'), 'TELEFON'),
    # Phone: landline with dash separators (e.g. "12-654-33-21", "71-344-89-01")
    (re.compile(r'(?<!\d)\d{2}-\d{3}-\d{2}-\d{2}(?!\d)'), 'TELEFON'),
    # Email
    (re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'), 'EMAIL'),
    # Polish postal code + city (e.g. "00-001 Warszawa")
    (re.compile(r'\d{2}-\d{3}\s+[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+'), 'ADRES'),
    # Street address (ul./al./os./pl. + name + optional number) — limit to max 4 words
    (re.compile(r'(?:ul\.|al\.|os\.|pl\.|Al\.)\s+[A-ZĄ-Ż][a-ząćęłńóśźż]+'
                r'(?:\s[A-ZĄ-Ż][a-ząćęłńóśźż]+){0,3}'
                r'(?:\s+\d+[a-zA-Z]?(?:/\d+[a-zA-Z]?)?)'), 'ADRES'),
    # NIP: 10 digits with dashes (e.g. "525-12-34-567" or "NIP: 5251234567")
    (re.compile(r'\b\d{3}-\d{2}-\d{2}-\d{3}\b'), 'NIP'),
    (re.compile(r'NIP[\s:]*\d{10}\b'), 'NIP'),
    # REGON: 9 or 14 digits (e.g. "REGON: 123456789")
    (re.compile(r'REGON[\s:]*\d{9}(?:\d{5})?\b'), 'REGON'),
    # KRS: 10 digits (e.g. "KRS: 0000234567" or "KRS 0000234567")
    (re.compile(r'KRS[\s:]*\d{10}\b'), 'KRS'),
    # PWZ (prawo wykonywania zawodu) — NOT anonymized, it's a public
    # professional license number for doctors, not personal data.
    # Dates: DD.MM.YYYY, DD-MM-YYYY, DD/MM/YYYY
    (re.compile(r'(?<!\d)\d{1,2}[./-]\d{1,2}[./-]\d{4}(?!\d)'), 'DATA'),
    # Written dates: "31 marca 2026" / "1 stycznia 2025 r."
    (re.compile(r'\d{1,2}\s+(?:' + _MONTHS_PL + r')\s+\d{4}(?:\s+r\.)?', re.I), 'DATA'),
]

# -- Name detection (regex-based, high recall for Polish documents) ---------

# Common Polish first names used as anchors for detecting name patterns
_PL_FIRST_NAMES = frozenset({
    'Adam', 'Adrian', 'Agata', 'Agnieszka', 'Aleksander', 'Aleksandra',
    'Alfred', 'Alicja', 'Alina', 'Amelia', 'Anastazja', 'Andrzej', 'Anna',
    'Antoni', 'Antonina', 'Arkadiusz', 'Artur',
    'Barbara', 'Bartlomiej', 'Bartosz', 'Beata', 'Benedykt', 'Bernadeta',
    'Blanka', 'Bogdan', 'Bogdana', 'Bogumil', 'Bogumila', 'Boguslaw',
    'Boguslawa', 'Boleslawa', 'Bozena', 'Bronislaw', 'Bronislawa',
    'Celina', 'Cezary', 'Czeslaw', 'Czeslaw',
    'Damian', 'Daniel', 'Daniela', 'Danuta', 'Dariusz', 'Dawid', 'Dominik',
    'Dominika', 'Dorota',
    'Edmund', 'Edward', 'Eleonora', 'Elzbieta', 'Emil', 'Emilia', 'Eugenia',
    'Eugeniusz', 'Ewa', 'Ewelina',
    'Fabian', 'Filip', 'Franciszek', 'Fryderyk',
    'Gabriel', 'Gabriela', 'Genowefa', 'Gertruda', 'Grazyna', 'Grzegorz',
    'Gustaw',
    'Halina', 'Hanna', 'Helena', 'Henryk', 'Henryka', 'Hubert',
    'Ignacy', 'Igor', 'Ilona', 'Irena', 'Ireneusz', 'Iwona', 'Izabela',
    'Jacek', 'Jadwiga', 'Jakub', 'Jan', 'Janina', 'Janusz', 'Jaroslaw',
    'Jerzy', 'Joanna', 'Jolanta', 'Jozef', 'Jozefa', 'Julia', 'Julian',
    'Juliusz', 'Justyna',
    'Kamil', 'Kamila', 'Karol', 'Karolina', 'Katarzyna', 'Kazimiera',
    'Kazimierz', 'Klaudia', 'Konrad', 'Kornelia', 'Krystian', 'Krystyna',
    'Krzysztof',
    'Laura', 'Leon', 'Leonard', 'Leszek', 'Lidia', 'Lilian', 'Lucjan',
    'Lucyna', 'Ludmila', 'Ludwik', 'Luiza', 'Lukasz',
    'Maciej', 'Magdalena', 'Maja', 'Maksymilian', 'Malgorzata', 'Marcel',
    'Marcin', 'Marek', 'Maria', 'Marian', 'Marianna', 'Mariusz', 'Marlena',
    'Marta', 'Mateusz', 'Michal', 'Michalina', 'Mieczyslaw', 'Milena',
    'Miroslaw', 'Miroslawa', 'Monika',
    'Natalia', 'Natasza', 'Nikola', 'Nikolaj', 'Nina', 'Norbert',
    'Olga', 'Oliwia', 'Oskar',
    'Patrycja', 'Patryk', 'Paulina', 'Pawel', 'Piotr', 'Przemyslaw',
    'Radoslaw', 'Rafal', 'Regina', 'Renata', 'Robert', 'Roman', 'Rozalia',
    'Rudolf', 'Ryszard',
    'Sabina', 'Sandra', 'Sebastian', 'Stanislaw', 'Stanislawa', 'Stefan',
    'Stefania', 'Sylwester', 'Sylwia', 'Szymon',
    'Tadeusz', 'Tatiana', 'Teresa', 'Tomasz', 'Tymoteusz',
    'Urszula',
    'Waldemar', 'Walentyna', 'Wanda', 'Weronika', 'Wieslaw', 'Wieslawa',
    'Wiktoria', 'Wiktor', 'Witold', 'Wladyslaw', 'Wladyslawa',
    'Wojciech',
    'Zbigniew', 'Zdzislaw', 'Zenon', 'Zofia', 'Zygmunt', 'Zyta',
})

# Declined (accusative/genitive) forms of common Polish first names
# Mapping: declined form -> base form (for matching)
_PL_FIRST_NAMES_DECLINED = {}
for _name in _PL_FIRST_NAMES:
    if _name.endswith('a') and len(_name) > 3:
        # feminine -a -> -e (acc), -y/-i (gen)
        _stem = _name[:-1]
        for _suf in ('e', 'y', 'i'):
            _PL_FIRST_NAMES_DECLINED[_stem + _suf] = _name
    elif not _name.endswith('a') and len(_name) > 3:
        # masculine consonant endings: +a (gen), +owi (dat), +em (inst)
        _PL_FIRST_NAMES_DECLINED[_name + 'a'] = _name
        _PL_FIRST_NAMES_DECLINED[_name + 'owi'] = _name
        _PL_FIRST_NAMES_DECLINED[_name + 'em'] = _name
        # special cases: names ending in -ek drop e: Marek->Marka
        if _name.endswith('ek'):
            _stem = _name[:-2] + 'k'
            _PL_FIRST_NAMES_DECLINED[_stem + 'a'] = _name
            _PL_FIRST_NAMES_DECLINED[_stem + 'owi'] = _name
            _PL_FIRST_NAMES_DECLINED[_stem + 'iem'] = _name
        # names ending in -sz: Tomasz->Tomasza, -usz: Tadeusz->Tadeusza
        if _name.endswith('sz') or _name.endswith('rz'):
            _PL_FIRST_NAMES_DECLINED[_name + 'a'] = _name

# Also handle declined surname patterns: -skiego/-skim/-skiemu, -ckiego/-ckim
_SURNAME_RE_DECLINED = (
    r'[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]{2,}'
    r'(?:-[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]{2,})?'
)

# Title prefixes that signal a person name follows
_TITLE_PREFIXES = (
    r'dr\s+n\.\s*med\.\s*',
    r'dr\s+hab\.\s*n\.\s*med\.\s*',
    r'prof\.\s*dr\s+hab\.\s*n\.\s*med\.\s*',
    r'prof\.\s*dr\s+hab\.\s*',
    r'prof\.\s*',
    r'dr\s+hab\.\s*',
    r'dr\s+',
    r'lek\.\s*med\.\s*',
    r'lek\.\s*',
    r'mgr\s+',
    r'inz\.\s*',
)

# Capitalized Polish surname (including compound): e.g. "Kowalski", "Kowalska-Nowak"
_SURNAME_RE = r'[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]{2,}(?:-[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]{2,})?'

# Polish name stopwords: words that look like names but aren't
_NAME_STOPWORDS = frozenset({
    'Pacjent', 'Pacjentka', 'Pacjenta', 'Lekarz', 'Konsultanci', 'Klinika',
    'Szpital', 'Centrum', 'Instytut', 'Poradnia', 'Oddzial', 'Pani', 'Pan',
    'Placowka', 'Adres', 'Telefon', 'Email', 'Podpis', 'Data', 'Numer',
    'Rozpoznanie', 'Badanie', 'Wyniki', 'Epikryza', 'Wnioski', 'Opinia',
    'Przebieg', 'Zalecenia', 'Kontrola', 'Osoba', 'Siostra', 'Brat',
    'Matka', 'Ojciec', 'Zona', 'Maz',
})


def _is_known_first_name(name):
    """Check if name is a known Polish first name (including declined forms)."""
    return name in _PL_FIRST_NAMES or name in _PL_FIRST_NAMES_DECLINED


def _surname_variants(surname):
    """Generate nominative + common declined forms of a Polish surname.

    Handles compound surnames (e.g. Kowalska-Nowak) by declining each part.
    """
    # For compound surnames, decline each part separately and combine
    if '-' in surname:
        parts = surname.split('-')
        all_forms = {surname}
        part_variants = [_surname_variants(p) for p in parts]
        # Add individual part variants
        for pv in part_variants:
            all_forms.update(pv)
        # Generate compound declined forms (e.g. Kowalskiej-Nowak)
        if len(part_variants) == 2:
            for v1 in part_variants[0]:
                for v2 in part_variants[1]:
                    all_forms.add(f"{v1}-{v2}")
        return all_forms

    forms = {surname}
    # Normalize to nominative if a declined form was passed in
    base = surname
    if surname.endswith('skiego'):
        base = surname[:-4] + 'i'       # Kowalskiego -> Kowalski
    elif surname.endswith('ckiego'):
        base = surname[:-4] + 'i'       # Nowickiego -> Nowicki
    elif surname.endswith('skiej'):
        base = surname[:-2] + 'a'       # Kowalskiej -> Kowalska
    elif surname.endswith('ckiej'):
        base = surname[:-2] + 'a'       # Nowickiej -> Nowicka
    elif surname.endswith('skiemu'):
        base = surname[:-3]             # Kowalskiemu -> Kowalski
    elif surname.endswith('ckiemu'):
        base = surname[:-3]             # Nowickiemu -> Nowicki
    elif surname.endswith('skim') and not surname.endswith('askim'):
        base = surname[:-1] + 'i'       # Kowalskim -> Kowalski
    elif surname.endswith('ckim'):
        base = surname[:-1] + 'i'       # Nowickim -> Nowicki
    elif surname.endswith('dzkiego'):
        base = surname[:-4] + 'i'       # Łódzkiego -> Łódzki
    elif surname.endswith('dzkim'):
        base = surname[:-1] + 'i'       # Łódzkim -> Łódzki

    forms.add(base)

    # Generate declined forms from base
    if base.endswith('ski'):
        stem = base[:-1]  # -ski -> -sk
        forms.update([stem + 'iego', stem + 'iemu', stem + 'im', stem + 'i'])
    elif base.endswith('cki'):
        stem = base[:-1]  # -cki -> -ck
        forms.update([stem + 'iego', stem + 'iemu', stem + 'im', stem + 'i'])
    elif base.endswith('dzki'):
        stem = base[:-1]  # -dzki -> -dzk
        forms.update([stem + 'iego', stem + 'iemu', stem + 'im', stem + 'i'])

    if base.endswith('ska'):
        stem = base[:-1]  # -ska -> -sk
        forms.update([stem + 'iej', stem + 'ą'])
    elif base.endswith('cka'):
        stem = base[:-1]  # -cka -> -ck
        forms.update([stem + 'iej', stem + 'ą'])
    elif base.endswith('a') and not base.endswith(('ska', 'cka')):
        # generic feminine: -a -> -ej, -ą
        forms.update([base[:-1] + 'ej', base[:-1] + 'ą'])
    elif not base.endswith('a'):
        # generic masculine consonant: +a (gen), +owi, +em
        forms.update([base + 'a', base + 'owi', base + 'em'])
        # Names ending in -ek drop the e: Dudek -> Dudka
        if base.endswith('ek'):
            stem = base[:-2] + 'k'
            forms.update([stem + 'a', stem + 'owi', stem + 'iem'])
        # Names ending in -ec: Kupiec -> Kupca
        if base.endswith('ec'):
            stem = base[:-2] + 'c'
            forms.update([stem + 'a', stem + 'owi', stem + 'em'])

    return forms


def _detect_names(text):
    """Detect Polish person names using pattern matching."""
    entities = []
    seen = set()
    known_surnames = set()  # collect detected surnames for cross-referencing

    # 1) Title + name patterns (dr, prof., lek. etc.) — search full text
    for prefix in _TITLE_PREFIXES:
        pat = re.compile(
            r'(?:' + prefix + r')'
            r'([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+'
            r'(?:[ \t]+[A-ZĄĆĘŁŃÓŚŹŻ]\.?[a-ząćęłńóśźż]*)?'
            r'[ \t]+' + _SURNAME_RE + r')'
        )
        for m in pat.finditer(text):
            name = m.group(1).strip()
            full = m.group(0).strip()
            if name not in seen and name.split()[0] not in _NAME_STOPWORDS:
                seen.add(name)
                seen.add(full)
                entities.append({'text': full, 'category': 'LEKARZ'})
                # extract surname for cross-ref
                parts = name.split()
                if parts:
                    for sp in parts[-1].split('-'):
                        if len(sp) >= 4:
                            known_surnames.add(sp)

    # 2) Known first name + surname(s) — search line-by-line
    name_pat = re.compile(
        r'\b([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)'
        r'(?:[ \t]+([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+))?'
        r'[ \t]+(' + _SURNAME_RE + r')\b'
    )
    for line in text.split('\n'):
        for m in name_pat.finditer(line):
            first = m.group(1)
            middle = m.group(2) or ''
            surname = m.group(3)
            full_match = m.group(0).strip()

            if not _is_known_first_name(first) and not _is_known_first_name(middle):
                continue
            if first in _NAME_STOPWORDS:
                continue
            if len(surname) < 3:
                continue
            if full_match not in seen:
                seen.add(full_match)
                entities.append({'text': full_match, 'category': 'IMIE_NAZWISKO'})
                for sp in surname.split('-'):
                    if len(sp) >= 4:
                        known_surnames.add(sp)

    # 2b) Surname-first: SURNAME Firstname (common in Polish official forms)
    surname_first_pat = re.compile(
        r'\b(' + _SURNAME_RE + r')'
        r'[ \t]+([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)\b'
    )
    for line in text.split('\n'):
        for m in surname_first_pat.finditer(line):
            surname_cand = m.group(1)
            first_cand = m.group(2)
            full_match = m.group(0).strip()
            if full_match in seen:
                continue
            if not _is_known_first_name(first_cand):
                continue
            if surname_cand in _NAME_STOPWORDS or first_cand in _NAME_STOPWORDS:
                continue
            if len(surname_cand) < 3:
                continue
            seen.add(full_match)
            entities.append({'text': full_match, 'category': 'IMIE_NAZWISKO'})
            for sp in surname_cand.split('-'):
                if len(sp) >= 4:
                    known_surnames.add(sp)

    # 3) Contextual label patterns — "Imię i nazwisko: NAME", "Syn: NAME" etc.
    _LABEL_PATTERNS = [
        r'[Ii]mi[eę]\s+i\s+nazwisko\s*:\s*',
        r'[Nn]azwisko\s+panie[nń]skie\s*:\s*',
        r'[Nn]azwisko\s*:\s*',
        r'[Ss]yn\s*:\s*',
        r'[Cc][oó]rka\s*:\s*',
        r'[Mm][aą][zż]\s*:\s*',
        r'[Żż]ona\s*:\s*',
        r'[Oo]jciec\s*:\s*',
        r'[Mm]atka\s*:\s*',
        r'[Oo]piekun(?:\s+prawny)?\s*:\s*',
        r'[Pp]rzedstawiciel(?:\s+ustawowy)?\s*(?:\([^)]*\))?\s*:\s*',
        r'[Oo]soba\s+upowa[zż]niona[^:]*:\s*',
        r'[Pp]e[lł]nomocnik[^:]*:\s*',
        r'[Bb]abcia\s+(?:macierzysta|ojczysta)\s*:\s*',
        r'[Dd]ziadek\s+(?:macierzysty|ojczysty)\s*:\s*',
    ]
    for label_re in _LABEL_PATTERNS:
        label_pat = re.compile(
            label_re + r'(' + _SURNAME_RE
            + r'(?:[ \t]+[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)*'
            + r'(?:[ \t]+' + _SURNAME_RE + r')?'
            + r')'
        )
        for m in label_pat.finditer(text):
            val = m.group(1).strip()
            if val and len(val) >= 3 and val not in seen:
                seen.add(val)
                entities.append({'text': val, 'category': 'IMIE_NAZWISKO'})
                for part in val.split():
                    for sp in part.split('-'):
                        if len(sp) >= 4 and sp[0].isupper():
                            known_surnames.add(sp)

    # 4) "z d." (z domu = maiden name) pattern
    maiden_pat = re.compile(
        r'z\s+d(?:omu)?\.\s+(' + _SURNAME_RE + r')'
    )
    for m in maiden_pat.finditer(text):
        maiden = m.group(0).strip()
        surname = m.group(1)
        if maiden not in seen:
            seen.add(maiden)
            entities.append({'text': maiden, 'category': 'IMIE_NAZWISKO'})
            for sp in surname.split('-'):
                if len(sp) >= 4:
                    known_surnames.add(sp)

    # 5) "K. Surname" abbreviation patterns
    abbrev_pat = re.compile(
        r'\b([A-ZĄĆĘŁŃÓŚŹŻ]\.)[ \t]+(' + _SURNAME_RE + r')\b'
    )
    for m in abbrev_pat.finditer(text):
        abbrev = m.group(0).strip()
        surname = m.group(2)
        if len(surname) >= 4 and abbrev not in seen:
            seen.add(abbrev)
            entities.append({'text': abbrev, 'category': 'IMIE_NAZWISKO'})
            for sp in surname.split('-'):
                if len(sp) >= 4:
                    known_surnames.add(sp)

    # 6) Surname cross-referencing: find standalone mentions of known surnames
    #    Expand each known surname into all declined forms
    if known_surnames:
        all_variants = set()
        for sn in known_surnames:
            all_variants.update(_surname_variants(sn))
        # filter out very short forms that could cause false positives
        all_variants = {v for v in all_variants if len(v) >= 4}
        surname_alt = '|'.join(re.escape(s) for s in sorted(all_variants, key=len, reverse=True))
        crossref_pat = re.compile(r'\b(' + surname_alt + r')\b')
        for m in crossref_pat.finditer(text):
            sname = m.group(1)
            if sname not in seen:
                start = m.start()
                if start > 0 and text[start-1].isalpha():
                    continue
                seen.add(sname)
                entities.append({'text': sname, 'category': 'IMIE_NAZWISKO'})

    return entities


def _detect_facilities(text):
    """Detect Polish medical facility names using pattern matching."""
    entities = []
    seen = set()

    # Use [^\n]* to stay within single lines
    facility_patterns = [
        re.compile(r'Szpital(?:u|em)?[ \t]+[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+'
                   r'(?:[ \t]+(?:im\.\s+)?[a-ząćęłńóśźżA-ZĄĆĘŁŃÓŚŹŻ.]+)*'),
        re.compile(r'Klinik[aięy][ \t]+[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+'
                   r'(?:[ \t]+[a-ząćęłńóśźżA-ZĄĆĘŁŃÓŚŹŻ]+){0,5}'),
        re.compile(r'Centrum[ \t]+(?:Medyczn[a-z]*|Zdrowia)[ \t]+[A-Za-ząćęłńóśźż]+'
                   r'(?:[ \t]+[A-Za-ząćęłńóśźż.]+)*'),
        re.compile(r'Poradni[aęy][ \t]+[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+'
                   r'(?:[ \t]+[a-ząćęłńóśźżA-ZĄĆĘŁŃÓŚŹŻ]+){0,4}'),
        re.compile(r'Instytut(?:u)?[ \t]+[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+'
                   r'(?:[ \t]+[a-ząćęłńóśźżA-ZĄĆĘŁŃÓŚŹŻ]+){0,5}'),
        re.compile(r'(?:Osrodek|Ośrodek)[ \t]+[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+'
                   r'(?:[ \t]+[a-ząćęłńóśźżA-ZĄĆĘŁŃÓŚŹŻ]+){0,3}'),
        re.compile(r'(?:NZOZ|SPZOZ|ZOZ)[ \t]+[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+'
                   r'(?:[ \t]+[a-ząćęłńóśźżA-ZĄĆĘŁŃÓŚŹŻ]+){0,3}'),
    ]

    for pat in facility_patterns:
        for m in pat.finditer(text):
            name = m.group(0).strip()
            # Trim trailing prepositions, articles, and title prefixes
            name = re.sub(r'[ \t]+(?:w|we|z|ze|na|przy|do|i|lub|oraz|dr|prof|lek|mgr)\.?[ \t]*$', '', name)
            if name not in seen and len(name) > 8:
                seen.add(name)
                entities.append({'text': name, 'category': 'NAZWA_PLACOWKI'})

    return entities


def _regex_detect(text):
    """Detect PII using regex patterns. Returns list of entity dicts."""
    # Normalize: collapse whitespace/newlines between digits
    normalized = re.sub(r'(\d)[\s\n]+(\d)', r'\1 \2', text)

    entities = []
    seen = set()
    for pattern, category in _REGEX_PATTERNS:
        for m in pattern.finditer(normalized):
            matched = m.group(0).strip()
            if not matched or len(matched) < 3:
                continue
            # Skip matches containing newlines (broken table fragments)
            if '\n' in matched:
                continue
            if matched not in seen:
                seen.add(matched)
                entities.append({'text': matched, 'category': category})

    # Remove entities that are substrings of a longer entity in same category
    final = []
    for e in entities:
        is_substring = False
        for other in entities:
            if other is not e and e['text'] in other['text'] and e['category'] == other['category']:
                is_substring = True
                break
        if not is_substring:
            final.append(e)
    return final


# -- LLM anonymization (for names, doctor names, facility names) ------------


def _call_llm(text_chunk):
    """Send text to the local LLM in a subprocess to avoid blocking gevent.

    Uses a result file instead of stdout because llama.cpp's C runtime
    writes CUDA/GGML init messages to stdout, polluting the JSON output.
    """
    import subprocess as _sp
    import tempfile

    # Write text to a temp file for the subprocess
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False,
                                     dir=_JOBS_DIR) as tf:
        tf.write(text_chunk[:3000])
        text_path = tf.name

    result_path = text_path + '.result.json'

    script = r'''
import sys, json, os, re
sys.path.insert(0, '/opt/ethos/backend')
sys.path.insert(0, '/opt/ethos/backend/blueprints')

text_path = sys.argv[1]
result_path = sys.argv[2]

with open(text_path) as f:
    text = f.read()

from model_library import get_library
lib = get_library()

# Prefer the best available Bielik model (Q8 > Q4) for Polish PII detection
bielik_preference = ['bielik-7b-q8', 'bielik-7b-q4']
best_bielik = None
downloaded = lib._config.get('downloaded', {})
for bid in bielik_preference:
    dl = downloaded.get(bid)
    if dl and os.path.isfile(dl.get('path', '')):
        best_bielik = bid
        break

if best_bielik:
    llm, err = lib.load_model(best_bielik)
else:
    llm, err = lib.load_model()

if err:
    with open(result_path, 'w') as rf:
        json.dump([], rf)
    sys.exit(0)
lib.touch_model()

system_prompt = (
    "Wypisz TYLKO imiona i nazwiska osob oraz nazwy placowek medycznych z tekstu.\n"
    "NIE wypisuj: rozpoznan (ICD-10), lekow, dawek, zalecen, dat, adresow, "
    "numerow PESEL/telefon/konta, email, wynikow badan.\n"
    "Kazda pozycja to KROTKI tekst (imie+nazwisko lub nazwa placowki) - "
    "max kilka slow, nigdy cale zdanie.\n"
    "Format odpowiedzi - TYLKO JSON tablica:\n"
    '[{"text":"Katarzyna Nowak","category":"IMIE_NAZWISKO"},'
    '{"text":"dr Jan Kowalski","category":"LEKARZ"},'
    '{"text":"Szpital Miejski","category":"NAZWA_PLACOWKI"}]\n'
    "Dozwolone kategorie: IMIE_NAZWISKO, LEKARZ, NAZWA_PLACOWKI.\n"
    "Bez komentarzy, bez markdown."
)
user_prompt = "Tekst:\n" + text + "\n\nJSON:"

resp = llm.create_chat_completion(
    messages=[
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ],
    max_tokens=1024,
    temperature=0.1,
)
content = resp["choices"][0]["message"]["content"].strip()
if "```" in content:
    m = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', content, re.DOTALL)
    if m:
        content = m.group(1).strip()

# Try to salvage truncated JSON arrays by closing them
def _try_parse_json_array(s):
    try:
        arr = json.loads(s)
        if isinstance(arr, dict) and "text" in arr:
            return [arr]
        if isinstance(arr, list):
            return arr
        return []
    except json.JSONDecodeError:
        pass
    # Truncated array: try closing the last complete element
    # Find last complete }, then close the array
    idx = s.rfind('}')
    if idx > 0:
        candidate = s[:idx+1] + ']'
        try:
            arr = json.loads(candidate)
            if isinstance(arr, list):
                return arr
        except json.JSONDecodeError:
            pass
    # Find embedded array
    m2 = re.search(r'\[.*\}', s, re.DOTALL)
    if m2:
        try:
            return json.loads(m2.group(0) + ']')
        except json.JSONDecodeError:
            pass
    return []

entities = _try_parse_json_array(content)
result = [e for e in entities if isinstance(e, dict) and "text" in e and "category" in e]
with open(result_path, 'w') as rf:
    json.dump(result, rf, ensure_ascii=False)
'''

    try:
        log.info('[doc_anonymizer] Starting LLM subprocess (text=%d chars)', len(text_chunk))
        proc = _sp.run(
            [sys.executable, '-c', script, text_path, result_path],
            capture_output=True, text=True,
            timeout=600,  # 10 min max
            cwd='/opt/ethos/backend',
        )
        log.info('[doc_anonymizer] LLM subprocess finished (rc=%s, stdout=%d, stderr=%d)',
                 proc.returncode, len(proc.stdout), len(proc.stderr))
        if proc.stderr:
            log.debug('[doc_anonymizer] LLM stderr: %s', proc.stderr[:300])
        if proc.returncode != 0:
            log.warning('[doc_anonymizer] LLM subprocess failed (rc=%s): %s',
                        proc.returncode, proc.stderr[:500])
            return []
        if not os.path.exists(result_path):
            log.warning('[doc_anonymizer] LLM result file not created. stdout=%s',
                        proc.stdout[:300])
            return []
        with open(result_path) as rf:
            result = json.load(rf)
        log.info('[doc_anonymizer] LLM returned %d entities', len(result))
        return result
    except _sp.TimeoutExpired:
        log.warning('[doc_anonymizer] LLM subprocess timed out (600s)')
        return []
    except Exception as e:
        log.warning('[doc_anonymizer] LLM subprocess error: %s', e)
        return []
    finally:
        for p in (text_path, result_path):
            if os.path.exists(p):
                os.unlink(p)


def _normalize_category(cat):
    """Strip diacritics and uppercase: IMIĘ -> IMIE, Nazwisko -> NAZWISKO."""
    nfkd = unicodedata.normalize('NFKD', cat)
    ascii_str = ''.join(c for c in nfkd if not unicodedata.combining(c))
    return ascii_str.upper().strip()


# Valid categories that the LLM is allowed to return
_VALID_LLM_CATEGORIES = frozenset({
    'IMIE_NAZWISKO', 'LEKARZ', 'NAZWA_PLACOWKI',
})

# ICD-10 code pattern (e.g. "I10", "E11.9", "M54.5", "J18.0")
_ICD_CODE_RE = re.compile(r'\b[A-Z]\d{2}(?:\.\d{1,2})?\b')

# Medical terms that should NEVER be anonymized
_MEDICAL_STOPWORDS = frozenset({
    # General medical terms
    'rozpoznanie', 'epikryza', 'zalecenia', 'leczenie', 'badanie', 'wyniki',
    'dawkowanie', 'kontrola', 'skierowanie', 'zaswiadczenie', 'zaświadczenie',
    'pacjent', 'pacjentka', 'choroba', 'zapalenie', 'niedokrwienna',
    'nadcisnienie', 'cukrzyca', 'hipercholesterolemia', 'diagnostyka',
    'rehabilitacja', 'operacja', 'zabieg', 'terapia', 'recepta',
    'objaw', 'zespol', 'skala', 'test', 'wynik', 'morfologia',
    'hemoglobina', 'leukocyty', 'trombocyty', 'erytrocyty', 'kreatynina',
    'bilirubina', 'glukoza', 'cholesterol', 'triglicerydy',
    # Diseases and conditions
    'niewydolnosc', 'migotanie', 'zatorowosc', 'zawal', 'udar',
    'padaczka', 'epilepsja', 'miazdzyca', 'nowotwor', 'bialaczka',
    'marskosc', 'niedokrwienie', 'zwezenie', 'torbiel', 'polip',
    'remisja', 'nawrot', 'przerzut', 'arytmia', 'bradykardia',
    'tachykardia', 'osteoporoza', 'reumatoidalne',
    # Procedures
    'gastroskopia', 'kolonoskopia', 'ultrasonografia', 'tomografia',
    'rezonans', 'echokardiografia', 'koronarografia', 'endoskopia',
    'biopsja', 'laparoskopia', 'hemodializa', 'chemioterapia',
    'radioterapia', 'ablacja', 'angioplastyka', 'holter',
    # Common medications (most frequent in Polish medical docs)
    'amlodypina', 'metformina', 'atorwastatyna', 'ramipril', 'bisoprolol',
    'enalapryl', 'peryndopryl', 'walsartan', 'telmisartan', 'losartan',
    'indapamid', 'torasemid', 'furosemid', 'spironolakton',
    'hydrochlorotiazyd', 'klopidogrel', 'warfaryna', 'dabigatran',
    'rywaroksaban', 'apiksaban', 'digoksyna', 'amiodaron',
    'gliklazyd', 'empagliflozyna', 'dapagliflozyna', 'semaglutyd',
    'insulina', 'ibuprofen', 'diklofenak', 'ketoprofen', 'paracetamol',
    'tramadol', 'metamizol', 'amoksycylina', 'azytromycyna',
    'ciprofloksacyna', 'doksycyklina', 'salbutamol', 'budezonid',
    'montelukast', 'omeprazol', 'pantoprazol', 'lansoprazol',
    'escytalopram', 'sertralina', 'wenlafaksyna', 'mirtazapina',
    'olanzapina', 'kwetiapina', 'lewotyroksyna', 'prednizon',
    'deksametazon', 'heparyna', 'acetylosalicylowy',
    'tikagrelol', 'lacydypina', 'kandesartan', 'liraglutyd',
    'cefaleksyna', 'teofilina', 'esomeprazol', 'alprazolam', 'diazepam',
    'finasteryd', 'proscar',
    # Medical abbreviations commonly misidentified as names
    'triglicerydy', 'fizjoterapeuty', 'ordynator',
    # Professional license numbers (public, not PII)
    'pwz', 'prawo', 'wykonywania', 'zawodu',
    # Anatomy
    'serce', 'pluca', 'watroba', 'nerki', 'trzustka', 'jelito',
    'zoladek', 'mozg', 'kregowy', 'przedsionek', 'komora',
    'zastawka', 'aorta', 'tetnica', 'zyla',
})


def _validate_llm_entities(entities):
    """Filter LLM output to remove garbage, diagnoses, and too-long entries."""
    valid = []
    for e in entities:
        text = e.get('text', '').strip()
        cat = _normalize_category(e.get('category', ''))

        # Reject empty or very short
        if len(text) < 3:
            continue

        # Reject entities that are too long (real names are short)
        if len(text) > 80:
            log.debug('[doc_anonymizer] LLM entity too long (%d chars), skipping: %s...', len(text), text[:50])
            continue

        # Reject invalid categories — LLM may hallucinate RZPOZNANIE, LECZKA etc.
        if cat not in _VALID_LLM_CATEGORIES:
            log.debug('[doc_anonymizer] LLM invalid category %s, skipping: %s', cat, text[:50])
            continue

        # Reject if it contains ICD codes (diagnoses)
        if _ICD_CODE_RE.search(text):
            log.debug('[doc_anonymizer] LLM entity contains ICD code, skipping: %s', text[:50])
            continue

        # Reject if it contains medical stopwords
        text_lower = text.lower()
        if any(sw in text_lower for sw in _MEDICAL_STOPWORDS):
            log.debug('[doc_anonymizer] LLM entity contains medical term, skipping: %s', text[:50])
            continue

        # Reject entities that look like full sentences (contain verbs/punctuation patterns)
        if text.count(' ') > 8:
            log.debug('[doc_anonymizer] LLM entity has too many words, skipping: %s', text[:50])
            continue

        valid.append({'text': text, 'category': cat})

    return valid


# -- Replacement logic ------------------------------------------------------

_PLACEHOLDER_MAP = {
    'IMIE': '[IMIE]',
    'NAZWISKO': '[NAZWISKO]',
    'IMIE_NAZWISKO': '[OSOBA]',
    'PESEL': '[PESEL]',
    'DATA_URODZENIA': '[DATA]',
    'DATA': '[DATA]',
    'ADRES': '[ADRES]',
    'TELEFON': '[TELEFON]',
    'EMAIL': '[EMAIL]',
    'NR_PACJENTA': '[NR_PACJENTA]',
    'NR_DOKUMENTU': '[NR_DOKUMENTU]',
    'NR_KONTA': '[NR_KONTA]',
    'NIP': '[NIP]',
    'REGON': '[REGON]',
    'KRS': '[KRS]',
    'NAZWA_PLACOWKI': '[PLACOWKA]',
    'LEKARZ': '[LEKARZ]',
    'INNE_PII': '[DANE_OSOBOWE]',
}


def _replace_entities_in_text(text, entities):
    """Replace PII entities with placeholders in text."""
    seen_texts = {}
    for e in entities:
        cat = _normalize_category(e.get('category', 'INNE_PII'))
        if e['text'] not in seen_texts:
            seen_texts[e['text']] = cat

    unique = [{'text': t, 'category': c} for t, c in seen_texts.items()]
    unique.sort(key=lambda e: len(e['text']), reverse=True)

    result = text
    replacements = []
    for e in unique:
        pii_text = e['text']
        category = e.get('category', 'INNE_PII')
        placeholder = _PLACEHOLDER_MAP.get(category, '[' + category + ']')

        if pii_text in result:
            count = result.count(pii_text)
            result = result.replace(pii_text, placeholder)
            replacements.append({
                'original': pii_text,
                'category': category,
                'placeholder': placeholder,
                'occurrences': count,
            })

    return result, replacements


# -- Document generation ----------------------------------------------------

def _redact_pdf(src_path, output_path, entities):
    """Redact PII in a PDF using blur overlay — preserves layout, hides text."""
    import fitz
    from io import BytesIO

    seen_texts = {}
    for e in entities:
        cat = _normalize_category(e.get('category', 'INNE_PII'))
        txt = e.get('text', '').strip()
        if txt and txt not in seen_texts:
            # Skip multi-line entities — they produce huge blur rects
            if '\n' in txt or '|' in txt:
                # Try to salvage by taking just the first meaningful segment
                parts = re.split(r'[|\n]+', txt)
                parts = [p.strip() for p in parts if len(p.strip()) >= 4]
                for p in parts:
                    if p not in seen_texts:
                        seen_texts[p] = (_PLACEHOLDER_MAP.get(cat, '[DANE]'), cat)
                continue
            seen_texts[txt] = (_PLACEHOLDER_MAP.get(cat, '[DANE]'), cat)

    sorted_items = sorted(seen_texts.items(), key=lambda x: len(x[0]), reverse=True)

    doc = fitz.open(src_path)
    total_redactions = []

    # Blur parameters
    blur_radius = 8
    try:
        from PIL import Image, ImageFilter
        has_pil = True
    except ImportError:
        has_pil = False
        log.warning('[doc_anonymizer] Pillow not installed, using solid redaction')

    for page_idx, page in enumerate(doc):
        page_rects = []  # (rect, placeholder, category)
        page_area = page.rect.width * page.rect.height
        max_rect_area = page_area * 0.25  # Skip rects > 25% of page

        for original, (placeholder, category) in sorted_items:
            instances = page.search_for(original)
            for inst in instances:
                rect_area = abs(inst.width * inst.height)
                if rect_area > max_rect_area:
                    log.debug('[doc_anonymizer] Skipping oversized rect (%.0f%% of page) for: %s',
                              rect_area / page_area * 100, original[:40])
                    continue
                page_rects.append((inst, placeholder, category, original))
                total_redactions.append({
                    'original': original,
                    'placeholder': placeholder,
                    'category': category,
                    'occurrences': 1,
                })

        if not page_rects:
            continue

        if has_pil:
            # Capture page pixmap BEFORE redaction for blur source
            zoom = 2  # 2x for quality
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Apply redactions with white fill to remove original text
            for rect, placeholder, category, original in page_rects:
                page.add_redact_annot(rect, text='', fill=(1, 1, 1))
            page.apply_redactions()

            # Insert blurred image patches and placeholder text
            for rect, placeholder, category, original in page_rects:
                # Scale rect coords to pixmap coords
                x0 = max(0, int(rect.x0 * zoom) - 2)
                y0 = max(0, int(rect.y0 * zoom) - 2)
                x1 = min(img.width, int(rect.x1 * zoom) + 2)
                y1 = min(img.height, int(rect.y1 * zoom) + 2)

                if x1 <= x0 or y1 <= y0:
                    continue

                # Crop, blur, save as PNG
                crop = img.crop((x0, y0, x1, y1))
                blurred = crop.filter(ImageFilter.GaussianBlur(radius=blur_radius))
                buf = BytesIO()
                blurred.save(buf, format='PNG')
                buf.seek(0)

                # Insert blurred image at original position
                page.insert_image(rect, stream=buf.getvalue())
        else:
            # Fallback: solid redaction (black fill, white text)
            for rect, placeholder, category, original in page_rects:
                page.add_redact_annot(
                    rect, text=placeholder, fontsize=0,
                    fill=(0, 0, 0), text_color=(1, 1, 1),
                )
            page.apply_redactions()

    # Add watermark on first page
    first = doc[0]
    first.insert_text(
        (first.rect.width - 180, 20),
        'DOKUMENT ZANONIMIZOWANY',
        fontsize=8, color=(0.5, 0.5, 0.5),
    )

    doc.save(output_path, garbage=4, deflate=True)
    doc.close()

    # Merge redaction counts
    seen = {}
    for r in total_redactions:
        key = r['original']
        if key not in seen:
            seen[key] = r
        else:
            seen[key]['occurrences'] += 1
    return list(seen.values())


def _is_scanned_pdf(filepath):
    """Check if a PDF is scanned (image-only, no text layer)."""
    try:
        import fitz
        doc = fitz.open(filepath)
        for page in doc:
            if page.get_text().strip():
                doc.close()
                return False
        doc.close()
        return True
    except Exception:
        return False


def _redact_scanned_pdf(src_path, output_path, entities):
    """Redact PII in a scanned PDF using OCR bounding boxes + blur."""
    import fitz
    from io import BytesIO
    try:
        import pytesseract
        from PIL import Image, ImageFilter
    except ImportError:
        log.error('[doc_anonymizer] OCR redaction requires pytesseract + Pillow')
        return []

    # Clean entities (same filter as text-based path)
    entities = _clean_entities(entities)

    # Build lookup of texts to redact
    seen_texts = {}
    for e in entities:
        cat = _normalize_category(e.get('category', 'INNE_PII'))
        txt = e.get('text', '').strip()
        if txt and txt not in seen_texts:
            seen_texts[txt] = cat

    if not seen_texts:
        # Nothing to redact, just copy
        import shutil
        shutil.copy2(src_path, output_path)
        return []

    # --- Categorize entities for targeted matching ---
    # Person-related categories: individual word matching is safe (names are unique)
    _PERSON_CATS = {'IMIE_NAZWISKO', 'LEKARZ', 'OSOBA'}
    # Numeric categories: match digit portions only
    _NUMERIC_CATS = {'PESEL', 'NIP', 'REGON', 'TELEFON', 'NUMER_KONTA'}
    # Date category: match exact string
    _DATE_CATS = {'DATA'}
    # Remaining (addresses, institutions, etc.): phrase-level matching only

    _SKIP_WORDS = {
        'lek', 'dr', 'med', 'prof', 'mgr', 'inz', 'hab', 'doc',
        'im', 'ul', 'al', 'os', 'pl', 'str', 'nr', 'tel', 'fax',
        'sp', 'zoo', 'nip', 'regon', 'krs', 'www', 'com', 'pwz',
    }

    # Build targeted lookups
    name_words = {}      # word -> [(entity_text, category)]
    numeric_patterns = []  # (digits_str, entity_text, category)
    date_strings = []    # (date_str, entity_text, category)
    phrase_entities = []  # (words_list, entity_text, category)

    for txt, cat in seen_texts.items():
        if cat in _PERSON_CATS:
            for word in txt.split():
                w = word.strip('.,;:!?()[]/-').lower()
                if len(w) >= 3 and w not in _SKIP_WORDS:
                    if w not in name_words:
                        name_words[w] = []
                    name_words[w].append((txt, cat))
        elif cat in _NUMERIC_CATS:
            digits = re.sub(r'[^\d]', '', txt)
            if len(digits) >= 7:
                numeric_patterns.append((digits, txt, cat))
        elif cat in _DATE_CATS:
            date_strings.append((txt.strip(), txt, cat))
        else:
            # Addresses, institutions — phrase matching
            words = [w.strip('.,;:!?()[]/-').lower()
                     for w in txt.split() if len(w.strip('.,;:!?()[]/-')) >= 3]
            words = [w for w in words if w not in _SKIP_WORDS]
            if words:
                phrase_entities.append((words, txt, cat))

    doc = fitz.open(src_path)
    total_redactions = []
    blur_radius = 12

    for page_idx, page in enumerate(doc):
        dpi = 300
        scale = dpi / 72
        mat = fitz.Matrix(scale, scale)
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        ocr_data = pytesseract.image_to_data(
            img, lang='pol+eng', config='--psm 6',
            output_type=pytesseract.Output.DICT
        )

        blur_rects = []
        matched_entities = set()
        n_words = len(ocr_data['text'])

        def _bbox(idx):
            return (ocr_data['left'][idx], ocr_data['top'][idx],
                    ocr_data['width'][idx], ocr_data['height'][idx])

        def _add_rect(idx, pad=4):
            x, y, w, h = _bbox(idx)
            blur_rects.append((
                max(0, x - pad), max(0, y - pad),
                min(img.width, x + w + pad), min(img.height, y + h + pad)
            ))

        def _add_span(start, count, pad=4):
            x0 = min(ocr_data['left'][start + j] for j in range(count))
            y0 = min(ocr_data['top'][start + j] for j in range(count))
            x1 = max(ocr_data['left'][start + j] + ocr_data['width'][start + j]
                     for j in range(count))
            y1 = max(ocr_data['top'][start + j] + ocr_data['height'][start + j]
                     for j in range(count))
            blur_rects.append((
                max(0, x0 - pad), max(0, y0 - pad),
                min(img.width, x1 + pad), min(img.height, y1 + pad)
            ))

        # 1) Person name words — match individual words
        for i in range(n_words):
            word = ocr_data['text'][i].strip()
            if not word:
                continue
            word_lower = word.strip('.,;:!?()[]/-').lower()
            if word_lower in name_words:
                _add_rect(i)
                for ent_txt, ent_cat in name_words[word_lower]:
                    matched_entities.add((ent_txt, ent_cat))

        # 2) Numeric IDs — match digit sequences against known IDs
        for i in range(n_words):
            word = ocr_data['text'][i].strip()
            if not word:
                continue
            word_digits = re.sub(r'[^\d]', '', word)
            if len(word_digits) < 7:
                continue
            for digits, ent_txt, ent_cat in numeric_patterns:
                if digits in word_digits or word_digits in digits:
                    _add_rect(i)
                    matched_entities.add((ent_txt, ent_cat))
                    break

        # 3) Dates — match exact date strings in OCR
        ocr_texts_lower = [ocr_data['text'][i].strip().lower() for i in range(n_words)]
        for date_str, ent_txt, ent_cat in date_strings:
            # Try single-word match (e.g. "10.10.2025")
            date_lower = date_str.lower()
            for i in range(n_words):
                if ocr_texts_lower[i] == date_lower:
                    _add_rect(i)
                    matched_entities.add((ent_txt, ent_cat))
            # Try two-word match (e.g. "10-10-" "2025")
            for i in range(n_words - 1):
                pair = ocr_texts_lower[i] + ocr_texts_lower[i + 1]
                pair_sp = ocr_texts_lower[i] + ' ' + ocr_texts_lower[i + 1]
                if date_lower == pair or date_lower == pair_sp:
                    _add_span(i, 2)
                    matched_entities.add((ent_txt, ent_cat))

        # 4) Phrases (addresses, institutions) — require majority of words to match
        # in a contiguous OCR window, but limit span to avoid full-page blur
        for phrase_words, ent_txt, ent_cat in phrase_entities:
            needed = max(2, len(phrase_words) * 2 // 3)  # at least 2/3 match
            pw_set = set(phrase_words)
            win = len(phrase_words) + 2  # allow a bit of slack
            for i in range(n_words - needed + 1):
                end = min(i + win, n_words)
                window_words = set()
                for j in range(i, end):
                    w = ocr_data['text'][j].strip('.,;:!?()[]/-').lower()
                    if w:
                        window_words.add(w)
                hits = pw_set & window_words
                if len(hits) >= needed:
                    # Check span doesn't cover too much of the page
                    x0 = min(ocr_data['left'][j] for j in range(i, end))
                    y0 = min(ocr_data['top'][j] for j in range(i, end))
                    x1 = max(ocr_data['left'][j]+ocr_data['width'][j] for j in range(i, end))
                    y1 = max(ocr_data['top'][j]+ocr_data['height'][j] for j in range(i, end))
                    span_area = (x1 - x0) * (y1 - y0)
                    page_area = img.width * img.height
                    if span_area > page_area * 0.10:
                        # Span too large — blur individual matching words instead
                        for j in range(i, end):
                            wl = ocr_data['text'][j].strip('.,;:!?()[]/-').lower()
                            if wl in pw_set:
                                _add_rect(j)
                    else:
                        _add_span(i, end - i)
                    matched_entities.add((ent_txt, ent_cat))
                    break

        if not blur_rects:
            continue

        # Apply blur to all matched rectangles, skip oversized ones
        page_area = img.width * img.height
        max_rect_area = page_area * 0.15
        for (x0, y0, x1, y1) in blur_rects:
            if x1 <= x0 or y1 <= y0:
                continue
            rect_area = (x1 - x0) * (y1 - y0)
            if rect_area > max_rect_area:
                log.debug('[doc_anonymizer] Skipping oversized scanned rect %.1f%% on page %d',
                          rect_area / page_area * 100, page_idx + 1)
                continue
            crop = img.crop((x0, y0, x1, y1))
            blurred = crop.filter(ImageFilter.GaussianBlur(radius=blur_radius))
            img.paste(blurred, (x0, y0))

        for ent_txt, ent_cat in matched_entities:
            total_redactions.append({
                'original': ent_txt,
                'placeholder': _PLACEHOLDER_MAP.get(ent_cat, '[DANE]'),
                'category': ent_cat,
                'occurrences': 1,
            })

        # Replace page image with blurred version (JPEG for smaller output)
        buf = BytesIO()
        img.save(buf, format='JPEG', quality=85, optimize=True)
        buf.seek(0)

        # Clear page and insert full blurred image
        page.clean_contents()
        page.insert_image(page.rect, stream=buf.getvalue())

    # Add watermark on first page
    first = doc[0]
    first.insert_text(
        (first.rect.width - 180, 20),
        'DOKUMENT ZANONIMIZOWANY',
        fontsize=8, color=(0.5, 0.5, 0.5),
    )

    doc.save(output_path, garbage=4, deflate=True)
    doc.close()

    # Merge redaction counts
    seen = {}
    for r in total_redactions:
        key = r['original']
        if key not in seen:
            seen[key] = r
        else:
            seen[key]['occurrences'] += 1
    return list(seen.values())


def _anonymize_docx_inplace(src_path, output_path, all_entities):
    """Anonymize a DOCX preserving original formatting."""
    import docx

    seen_texts = {}
    for e in all_entities:
        cat = _normalize_category(e.get('category', 'INNE_PII'))
        if e['text'] not in seen_texts:
            seen_texts[e['text']] = cat

    unique = [{'text': t, 'category': c} for t, c in seen_texts.items()]
    unique.sort(key=lambda e: len(e['text']), reverse=True)

    replacements = {}
    for e in unique:
        placeholder = _PLACEHOLDER_MAP.get(e['category'], '[' + e['category'] + ']')
        replacements[e['text']] = placeholder

    document = docx.Document(src_path)

    def _replace_in_paragraph(para):
        full_text = para.text
        if not full_text.strip():
            return
        new_text = full_text
        for original, placeholder in replacements.items():
            new_text = new_text.replace(original, placeholder)
        if new_text == full_text:
            return
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ''
        else:
            para.text = new_text

    for para in document.paragraphs:
        _replace_in_paragraph(para)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para)

    document.save(output_path)


# -- Background anonymization task -----------------------------------------

_active_jobs = {}
_jobs_lock = threading.Lock()

# Sequential job queue — process one document at a time to save resources
_job_queue = []       # list of (job_id, src_path, filename, ext, username)
_queue_lock = threading.Lock()
_worker_running = False


def _cleanup_stuck_jobs():
    """Mark any 'processing' jobs as 'error' on startup (server crashed)."""
    if not os.path.isdir(_JOBS_DIR):
        return
    for entry in os.listdir(_JOBS_DIR):
        meta_path = os.path.join(_JOBS_DIR, entry, 'meta.json')
        if not os.path.isfile(meta_path):
            continue
        try:
            with open(meta_path) as f:
                meta = json.load(f)
            if meta.get('status') == 'processing':
                meta['status'] = 'error'
                meta['error'] = 'Serwer zostal zrestartowany w trakcie przetwarzania.'
                with open(meta_path, 'w') as f:
                    json.dump(meta, f, ensure_ascii=False, indent=2)
                log.info('[doc_anonymizer] Marked stuck job %s as error', entry)
        except Exception:
            continue


_cleanup_stuck_jobs()


def _emit_progress(job_id, stage, percent, message):
    with _jobs_lock:
        if job_id in _active_jobs:
            _active_jobs[job_id].update({
                'stage': stage, 'progress': percent, 'message': message,
            })
    sio = getattr(doc_anonymizer_bp, '_socketio', None)
    if sio:
        sio.emit('anon_progress', {
            'job_id': job_id, 'stage': stage,
            'percent': percent, 'message': message,
        })


_GARBAGE_NAME_PATTERNS = re.compile(
    r'^(?:Oun|Opl|Nmr|LVH|RBBB|LBBB|AoVmax|AcT|PG|GLTW|D\.S\.|EF|'
    r'Fizjoterapeuty|Ordynator|Pielegniar|Rehabilitant|Technik|Lica|'
    r'V+i*\s*i?\s*V*l*|VII+|VIII?|Vlll|'
    r'bz|max|min|sp\.|Sp\.|Meen)$',
    re.IGNORECASE
)

_MEDICAL_ABBREV_STOPWORDS = frozenset({
    'lvh', 'rbbb', 'lbbb', 'af', 'ef', 'aovmax', 'act', 'pg',
    'gltw', 'nmr', 'mri', 'ct', 'usg', 'ekg', 'emg', 'eeg',
    'tsh', 'ft3', 'ft4', 'crp', 'opl', 'oun', 'ast', 'alt',
    'bnp', 'gfr', 'hba1c', 'ldl', 'hdl', 'wbc', 'rbc', 'plt',
    'hgb', 'mch', 'mchc', 'mcv', 'inr', 'aptt', 'd.s.', 'ds',
    'lica', 'meen', 'pwz',
})


def _clean_entities(entities):
    """Remove garbage entities before redaction."""
    cleaned = []
    for e in entities:
        txt = e.get('text', '').strip()
        cat = e.get('category', '')

        # Strip trailing pipe/special chars from PDF extraction artifacts
        txt = re.sub(r'[|¢©]+\s*$', '', txt).strip()
        txt = re.sub(r'^[|¢©]+\s*', '', txt).strip()
        if not txt:
            continue
        e = dict(e, text=txt)

        # Skip empty or very short non-numeric entities
        if len(txt) < 3 and cat in ('IMIE_NAZWISKO', 'IMIE', 'NAZWISKO', 'LEKARZ', 'NAZWA_PLACOWKI'):
            continue

        # Skip entities containing copyright symbols or URL patterns
        if '©' in txt or 'www.' in txt or '.com' in txt or 'http' in txt:
            continue

        # Skip known garbage patterns
        if _GARBAGE_NAME_PATTERNS.match(txt):
            continue

        # Skip medical abbreviations detected as names
        if txt.lower().replace('.', '').replace(' ', '') in _MEDICAL_ABBREV_STOPWORDS:
            continue

        # Skip name entities that START with a medical stopword
        if cat in ('IMIE_NAZWISKO', 'IMIE', 'NAZWISKO'):
            first_word = txt.split()[0].lower() if txt.split() else ''
            if first_word in _MEDICAL_STOPWORDS or first_word in _MEDICAL_ABBREV_STOPWORDS:
                continue

        # Skip entities that look like job titles, not names
        title_words = {'fizjoterapeuty', 'ordynator', 'pielęgniarka', 'pielegniar',
                       'rehabilitant', 'technik', 'dietetyk', 'logopeda', 'psycholog'}
        if txt.lower() in title_words:
            continue

        # Skip PWZ (prawo wykonywania zawodu) — public professional license, not PII
        if re.match(r'(?i)(?:nr\s+)?PWZ[\s:]*\d{5,7}', txt):
            continue
        if cat == 'NR_DOKUMENTU' and 'PWZ' in txt.upper():
            continue

        # Split multi-line/pipe entities into clean parts
        if '\n' in txt or '|' in txt:
            parts = re.split(r'[|\n]+', txt)
            good_parts = [p.strip() for p in parts
                          if len(p.strip()) >= 4 and '©' not in p and 'www.' not in p]
            for p in good_parts:
                cleaned.append({'text': p, 'category': cat})
            continue

        cleaned.append(e)

    log.info('[doc_anonymizer] Entity cleanup: %d -> %d entities', len(entities), len(cleaned))
    return cleaned


def _run_anonymization(job_id, src_path, filename, file_ext, username):
    """Background: extract text -> LLM -> replace -> generate output."""
    job = _job_dir(job_id)
    meta_path = os.path.join(job, 'meta.json')

    try:
        _emit_progress(job_id, 'extracting', 10,
                       'Wyodrebnianie tekstu z dokumentu...')

        if file_ext == '.pdf':
            text_parts = _extract_text_pdf(src_path)
        elif file_ext in ('.docx', '.doc'):
            text_parts = _extract_text_docx(src_path)
        else:
            raise ValueError('Unsupported file type: ' + file_ext)

        if not text_parts or all(not t.strip() for t in text_parts):
            raise ValueError(
                'Nie udalo sie wyodrebnic tekstu z dokumentu. '
                'Plik moze byc zeskanowany — OCR nie rozpoznal tekstu.')

        total_parts = len(text_parts)
        all_entities = []

        # For PDFs, normalize text to fix fragmented extraction
        # (digits split across lines, newlines in bank accounts, etc.)
        if file_ext == '.pdf':
            text_parts = [re.sub(r'(\d)[\s\n]+(\d)', r'\1 \2', p) for p in text_parts]

        # Concatenate all text for detection (catches items spanning pages)
        full_text = '\n\n'.join(t for t in text_parts if t.strip())
        # Normalize digit sequences in full text (catches cross-page items)
        full_text = re.sub(r'(\d)[\s\n]+(\d)', r'\1 \2', full_text)

        # Phase 1: Regex-based detection on full text (fast, reliable)
        _emit_progress(job_id, 'regex', 20,
                       'Wykrywanie PESEL, telefonow, adresow (regex)...')
        regex_hits = _regex_detect(full_text)
        all_entities.extend(regex_hits)

        # Phase 2: Name & facility detection (regex-based, high recall)
        _emit_progress(job_id, 'analyzing', 40,
                       'Wykrywanie imion, nazwisk i placowek...')
        name_hits = _detect_names(full_text)
        facility_hits = _detect_facilities(full_text)
        all_entities.extend(name_hits)
        all_entities.extend(facility_hits)

        # Phase 3: spaCy NER detection (fast, replaces LLM)
        _emit_progress(job_id, 'analyzing', 55,
                       'Analiza NER (imiona/nazwiska/placowki)...')
        try:
            ner_entities = _call_spacy_ner(full_text)
            # Only add NER entities not already found by regex/name detection
            existing_lower = {e['text'].lower() for e in all_entities}
            for ne in ner_entities:
                ne_lower = ne.get('text', '').lower()
                if ne_lower and ne_lower not in existing_lower:
                    # Also check if it's a substring of an existing entity
                    is_dup = any(ne_lower in ex for ex in existing_lower)
                    if not is_dup:
                        all_entities.append(ne)
                        existing_lower.add(ne_lower)
        except Exception as e:
            log.warning('[doc_anonymizer] spaCy NER error: %s', e)
            # Fallback to LLM if spaCy fails
            try:
                llm_entities = _call_llm(full_text)
                llm_entities = _validate_llm_entities(llm_entities)
                existing_lower = {e['text'].lower() for e in all_entities}
                for le in llm_entities:
                    if le.get('text', '').lower() not in existing_lower:
                        all_entities.append(le)
                        existing_lower.add(le['text'].lower())
            except Exception as e2:
                log.warning('[doc_anonymizer] LLM fallback also failed: %s', e2)

        _emit_progress(job_id, 'replacing', 85,
                       'Zastepowanie danych osobowych...')

        # Clean up entities — remove garbage before redaction
        all_entities = _clean_entities(all_entities)

        _emit_progress(job_id, 'generating', 90,
                       'Generowanie zanonimizowanego dokumentu...')

        out_filename = 'anonymized_' + filename
        out_path = os.path.join(job, out_filename)

        if file_ext == '.pdf':
            if _is_scanned_pdf(src_path):
                log.info('[doc_anonymizer] Scanned PDF detected, using OCR redaction')
                redact_repls = _redact_scanned_pdf(src_path, out_path, all_entities)
            else:
                redact_repls = _redact_pdf(src_path, out_path, all_entities)
            seen_repls = {}
            for r in redact_repls:
                key = r['original']
                if key not in seen_repls:
                    seen_repls[key] = r
                else:
                    seen_repls[key]['occurrences'] += r['occurrences']
        elif file_ext in ('.docx', '.doc'):
            _anonymize_docx_inplace(src_path, out_path, all_entities)
            seen_repls = {}
            for e in all_entities:
                txt = e.get('text', '').strip()
                if not txt or txt in seen_repls:
                    continue
                cat = _normalize_category(e.get('category', 'INNE_PII'))
                seen_repls[txt] = {
                    'original': txt,
                    'placeholder': _PLACEHOLDER_MAP.get(cat, '[DANE]'),
                    'category': cat,
                    'occurrences': 1,
                }
        else:
            seen_repls = {}

        meta = {
            'job_id': job_id,
            'status': 'done',
            'filename': filename,
            'output_filename': out_filename,
            'file_ext': file_ext,
            'username': username,
            'created_at': time.time(),
            'completed_at': time.time(),
            'entities_found': len(seen_repls),
            'replacements': list(seen_repls.values()),
            'pages_analyzed': total_parts,
        }
        with open(meta_path, 'w') as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)

        _emit_progress(job_id, 'done', 100,
                       'Zanonimizowano - znaleziono %d danych osobowych' % len(seen_repls))

        with _jobs_lock:
            if job_id in _active_jobs:
                _active_jobs[job_id]['status'] = 'done'

    except Exception as e:
        log.error('[doc_anonymizer] Anonymization failed for %s: %s', job_id, e)
        meta = {
            'job_id': job_id, 'status': 'error',
            'filename': filename, 'username': username,
            'created_at': time.time(), 'error': str(e),
        }
        with open(meta_path, 'w') as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)
        _emit_progress(job_id, 'error', 0, str(e))
        with _jobs_lock:
            if job_id in _active_jobs:
                _active_jobs[job_id]['status'] = 'error'


def _queue_worker():
    """Process queued jobs one at a time."""
    global _worker_running
    while True:
        with _queue_lock:
            if not _job_queue:
                _worker_running = False
                return
            job_args = _job_queue.pop(0)

        job_id = job_args[0]
        with _jobs_lock:
            if job_id in _active_jobs:
                _active_jobs[job_id]['status'] = 'processing'

        _run_anonymization(*job_args)


def _enqueue_job(job_id, src_path, filename, ext, username):
    """Add a job to the sequential queue and start worker if needed."""
    global _worker_running
    with _queue_lock:
        _job_queue.append((job_id, src_path, filename, ext, username))
        if not _worker_running:
            _worker_running = True
            t = threading.Thread(target=_queue_worker, daemon=True)
            t.start()


# -- Routes -----------------------------------------------------------------

@doc_anonymizer_bp.route('/upload', methods=['POST'])
def anon_upload():
    """Upload a PDF/DOCX file for anonymization."""
    _ensure_jobs_dir()
    _ensure_deps()

    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    f = request.files['file']
    if not f.filename:
        return jsonify({'error': 'Empty filename'}), 400

    filename = f.filename
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ('.pdf', '.docx', '.doc'):
        return jsonify({'error': 'Obslugiwane formaty: PDF, DOCX'}), 400

    # Check NER/AI capability — spaCy preferred, LLM as bonus
    has_spacy = _load_spacy_model() is not None
    has_llm = False
    try:
        from model_library import get_library
        lib = get_library()
        downloaded = lib._config.get('downloaded', {})
        has_llm = any(
            downloaded.get(bid) and os.path.isfile(downloaded[bid].get('path', ''))
            for bid in ('bielik-7b-q8', 'bielik-7b-q4')
        )
        if not has_llm:
            active = lib.get_active_model()
            has_llm = active is not None
    except ImportError:
        pass

    if not has_spacy and not has_llm:
        return jsonify({
            'error': 'Brak modelu NER ani LLM. '
                     'Uruchom ponownie serwer lub zainstaluj AI Assistant.'
        }), 400

    job_id = uuid.uuid4().hex[:12]
    job = _job_dir(job_id)
    os.makedirs(job, exist_ok=True)

    src_path = os.path.join(job, 'original' + ext)
    f.save(src_path)

    username = _get_username()

    meta = {
        'job_id': job_id, 'status': 'processing',
        'filename': filename, 'file_ext': ext,
        'username': username, 'created_at': time.time(),
    }
    with open(os.path.join(job, 'meta.json'), 'w') as mf:
        json.dump(meta, mf, ensure_ascii=False)

    with _jobs_lock:
        _active_jobs[job_id] = {'status': 'queued', 'progress': 0}

    _enqueue_job(job_id, src_path, filename, ext, username)

    return jsonify({'ok': True, 'job_id': job_id, 'filename': filename})


@doc_anonymizer_bp.route('/jobs', methods=['GET'])
def anon_jobs():
    """List all anonymization jobs for the current user, newest first."""
    _ensure_jobs_dir()
    username = _get_username()
    jobs = []

    for entry in os.listdir(_JOBS_DIR):
        meta_path = os.path.join(_JOBS_DIR, entry, 'meta.json')
        if not os.path.isfile(meta_path):
            continue
        try:
            with open(meta_path) as f:
                meta = json.load(f)
            if meta.get('username') == username or getattr(g, 'role', None) == 'admin':
                with _jobs_lock:
                    live = _active_jobs.get(entry, {})
                if live:
                    live_status = live.get('status', '')
                    if live_status == 'queued' and meta.get('status') == 'processing':
                        meta['status'] = 'queued'
                        meta['message'] = 'W kolejce...'
                    elif meta.get('status') == 'processing':
                        meta['progress'] = live.get('progress', 0)
                        meta['message'] = live.get('message', '')
                jobs.append(meta)
        except Exception:
            continue

    # Sort by created_at descending (newest first)
    jobs.sort(key=lambda j: j.get('created_at', 0), reverse=True)

    return jsonify({'items': jobs})


@doc_anonymizer_bp.route('/download/<job_id>', methods=['GET'])
def anon_download(job_id):
    """Download the anonymized file."""
    job = _job_dir(job_id)
    meta_path = os.path.join(job, 'meta.json')

    if not os.path.isfile(meta_path):
        return jsonify({'error': 'Job not found'}), 404

    with open(meta_path) as f:
        meta = json.load(f)

    if meta.get('status') != 'done':
        return jsonify({'error': 'Anonymization not complete'}), 400

    out_file = os.path.join(job, meta['output_filename'])
    if not os.path.isfile(out_file):
        return jsonify({'error': 'Output file missing'}), 404

    if meta['file_ext'] == '.pdf':
        mime = 'application/pdf'
    else:
        mime = ('application/vnd.openxmlformats-officedocument'
                '.wordprocessingml.document')

    return send_file(out_file, mimetype=mime, as_attachment=True,
                     download_name=meta['output_filename'])


@doc_anonymizer_bp.route('/job/<job_id>', methods=['DELETE'])
def anon_delete_job(job_id):
    """Delete an anonymization job and its files."""
    import shutil
    job = _job_dir(job_id)
    if os.path.isdir(job):
        shutil.rmtree(job, ignore_errors=True)
    with _jobs_lock:
        _active_jobs.pop(job_id, None)
    return jsonify({'ok': True})


@doc_anonymizer_bp.route('/jobs/delete-batch', methods=['POST'])
def anon_delete_batch():
    """Delete multiple anonymization jobs at once."""
    import shutil
    data = request.get_json(silent=True) or {}
    job_ids = data.get('job_ids', [])
    if not isinstance(job_ids, list) or not job_ids:
        return jsonify({'error': 'job_ids list required'}), 400
    deleted = 0
    for jid in job_ids:
        if not isinstance(jid, str) or not jid:
            continue
        job = _job_dir(jid)
        if os.path.isdir(job):
            shutil.rmtree(job, ignore_errors=True)
            deleted += 1
        with _jobs_lock:
            _active_jobs.pop(jid, None)
    return jsonify({'ok': True, 'deleted': deleted})


@doc_anonymizer_bp.route('/preview/<job_id>/<which>', methods=['GET'])
def anon_preview(job_id, which):
    """Serve original or anonymized file inline for preview.

    ``which`` must be 'original' or 'anonymized'.
    PDFs are served as raw files (for PDF.js rendering in the frontend).
    DOCX text is extracted and returned as JSON.
    """
    if which not in ('original', 'anonymized'):
        return jsonify({'error': 'Invalid preview type'}), 400

    job = _job_dir(job_id)
    meta_path = os.path.join(job, 'meta.json')
    if not os.path.isfile(meta_path):
        return jsonify({'error': 'Job not found'}), 404

    with open(meta_path) as f:
        meta = json.load(f)

    file_ext = meta.get('file_ext', '.pdf')

    if which == 'original':
        file_path = os.path.join(job, 'original' + file_ext)
    else:
        if meta.get('status') != 'done':
            return jsonify({'error': 'Anonymization not complete'}), 400
        file_path = os.path.join(job, meta.get('output_filename', ''))

    if not os.path.isfile(file_path):
        return jsonify({'error': 'File not found'}), 404

    # PDF: serve raw file for PDF.js rendering
    if file_ext == '.pdf':
        return send_file(file_path, mimetype='application/pdf',
                         as_attachment=False)

    # DOCX: extract text and return as JSON
    try:
        paragraphs = _extract_text_docx(file_path)
        return jsonify({'ok': True, 'text': '\n\n'.join(paragraphs)})
    except Exception as e:
        log.warning('[doc_anonymizer] Preview text extraction failed: %s', e)
        return jsonify({'error': 'Could not extract text: ' + str(e)}), 500


@doc_anonymizer_bp.route('/regenerate/<job_id>', methods=['POST'])
def anon_regenerate(job_id):
    """Re-generate anonymized document excluding specified replacements.

    Accepts JSON ``{"excluded": ["original text 1", "original text 2"]}``.
    Items in *excluded* are treated as false positives — they will NOT be
    redacted.  The false-positive feedback is stored for future model
    retraining.
    """
    job = _job_dir(job_id)
    meta_path = os.path.join(job, 'meta.json')
    if not os.path.isfile(meta_path):
        return jsonify({'error': 'Job not found'}), 404

    with open(meta_path) as f:
        meta = json.load(f)

    if meta.get('status') != 'done':
        return jsonify({'error': 'Job not finished yet'}), 400

    data = request.get_json(silent=True) or {}
    if isinstance(data, str):
        try:
            data = json.loads(data)
        except Exception:
            data = {}
    excluded = set(data.get('excluded', []))
    if not excluded:
        return jsonify({'error': 'Nothing to exclude'}), 400

    file_ext = meta.get('file_ext', '.pdf')
    src_path = os.path.join(job, 'original' + file_ext)
    if not os.path.isfile(src_path):
        return jsonify({'error': 'Original file missing'}), 404

    # Save false-positive feedback for model retraining
    fb_path = os.path.join(job, 'feedback.json')
    feedback = []
    if os.path.isfile(fb_path):
        try:
            with open(fb_path) as f:
                feedback = json.load(f)
        except Exception:
            feedback = []
    for orig in excluded:
        matching = [r for r in meta.get('replacements', [])
                    if r.get('original') == orig]
        for r in matching:
            fb_entry = {
                'original': orig,
                'category': r.get('category', ''),
                'action': 'false_positive',
                'timestamp': time.time(),
            }
            if fb_entry not in feedback:
                feedback.append(fb_entry)
    with open(fb_path, 'w') as f:
        json.dump(feedback, f, ensure_ascii=False, indent=2)

    # Also save to global feedback file for model retraining
    global_fb_path = os.path.join(os.path.dirname(job), 'false_positives.json')
    global_fb = []
    if os.path.isfile(global_fb_path):
        try:
            with open(global_fb_path) as f:
                global_fb = json.load(f)
        except Exception:
            global_fb = []
    for orig in excluded:
        matching = [r for r in meta.get('replacements', [])
                    if r.get('original') == orig]
        for r in matching:
            global_fb.append({
                'original': orig,
                'category': r.get('category', ''),
                'job_id': job_id,
                'filename': meta.get('filename', ''),
                'timestamp': time.time(),
            })
    with open(global_fb_path, 'w') as f:
        json.dump(global_fb, f, ensure_ascii=False, indent=2)

    # Rebuild entity list from original replacements, minus excluded
    all_repls = meta.get('replacements', [])
    kept_entities = []
    for r in all_repls:
        if r.get('original') not in excluded:
            kept_entities.append({
                'text': r['original'],
                'category': r.get('category', 'INNE_PII'),
            })

    # Re-generate the output document
    out_filename = meta.get('output_filename', 'anonymized_' + meta.get('filename', 'doc'))
    out_path = os.path.join(job, out_filename)

    try:
        if file_ext == '.pdf':
            if _is_scanned_pdf(src_path):
                redact_repls = _redact_scanned_pdf(src_path, out_path, kept_entities)
            else:
                redact_repls = _redact_pdf(src_path, out_path, kept_entities)
            seen_repls = {}
            for r in redact_repls:
                key = r['original']
                if key not in seen_repls:
                    seen_repls[key] = r
                else:
                    seen_repls[key]['occurrences'] += r['occurrences']
        elif file_ext in ('.docx', '.doc'):
            _anonymize_docx_inplace(src_path, out_path, kept_entities)
            seen_repls = {}
            for e in kept_entities:
                txt = e.get('text', '').strip()
                if not txt or txt in seen_repls:
                    continue
                cat = _normalize_category(e.get('category', 'INNE_PII'))
                seen_repls[txt] = {
                    'original': txt,
                    'placeholder': _PLACEHOLDER_MAP.get(cat, '[DANE]'),
                    'category': cat,
                    'occurrences': 1,
                }
        else:
            return jsonify({'error': 'Unsupported format'}), 400

        # Update meta
        meta['replacements'] = list(seen_repls.values())
        meta['entities_found'] = len(seen_repls)
        meta['excluded'] = list(excluded)
        meta['regenerated_at'] = time.time()
        with open(meta_path, 'w') as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)

        return jsonify({
            'ok': True,
            'entities_found': len(seen_repls),
            'excluded_count': len(excluded),
            'replacements': list(seen_repls.values()),
        })
    except Exception as e:
        log.error('[doc_anonymizer] Regeneration failed for %s: %s', job_id, e)
        return jsonify({'error': 'Regeneration failed: ' + str(e)}), 500


@doc_anonymizer_bp.route('/status', methods=['GET'])
def anon_app_status():
    """Check app status and AI Chat dependency."""
    import shutil

    has_fitz = True
    try:
        import fitz  # noqa: F401
    except ImportError:
        has_fitz = False

    has_pypdf2 = True
    try:
        import PyPDF2  # noqa: F401
    except ImportError:
        has_pypdf2 = False

    has_pillow = True
    try:
        from PIL import Image  # noqa: F401
    except ImportError:
        has_pillow = False

    has_tesseract = bool(shutil.which('tesseract'))
    has_pytesseract = True
    try:
        import pytesseract  # noqa: F401
    except ImportError:
        has_pytesseract = False

    # Check spaCy model availability
    has_spacy = False
    spacy_model_name = None
    pii_model_path = os.path.join(data_path('models'), 'spacy_pii_pl')
    if os.path.isdir(pii_model_path) and os.path.isfile(os.path.join(pii_model_path, 'meta.json')):
        has_spacy = True
        spacy_model_name = 'spacy_pii_pl (fine-tuned)'
    else:
        try:
            import spacy
            spacy.load('pl_core_news_lg')
            has_spacy = True
            spacy_model_name = 'pl_core_news_lg (base)'
        except Exception:
            pass

    result = {
        'installed': True,
        'ai_chat_available': False,
        'model_loaded': False,
        'active_model': None,
        'recommended_model': 'spaCy NER (wbudowany)',
        'deps_ok': has_fitz and has_pypdf2 and has_pillow and bool(shutil.which('pdftotext')),
        'has_fitz': has_fitz,
        'has_pypdf2': has_pypdf2,
        'has_pillow': has_pillow,
        'has_pdftotext': bool(shutil.which('pdftotext')),
        'has_ocr': has_tesseract and has_pytesseract,
        'has_tesseract': has_tesseract,
        'has_pytesseract': has_pytesseract,
        'has_spacy': has_spacy,
        'spacy_model': spacy_model_name,
    }

    try:
        from model_library import get_library
        result['ai_chat_available'] = True
        lib = get_library()
        active = lib.get_active_model()
        if active:
            result['active_model'] = active.get('name', active.get('id'))
        loaded_llm, loaded_id = lib.get_loaded_model()
        result['model_loaded'] = loaded_llm is not None

        # Check for downloaded Bielik models (preferred for anonymization)
        downloaded = lib._config.get('downloaded', {})
        for bid in ('bielik-7b-q8', 'bielik-7b-q4'):
            dl = downloaded.get(bid)
            if dl and os.path.isfile(dl.get('path', '')):
                result['bielik_model'] = bid
                break
    except ImportError:
        pass

    return jsonify(result)


# -- Package routes (install / uninstall / pkg-status) ----------------------

def _on_uninstall(wipe):
    if wipe:
        import shutil
        if os.path.isdir(_JOBS_DIR):
            shutil.rmtree(_JOBS_DIR, ignore_errors=True)


register_pkg_routes(
    doc_anonymizer_bp,
    install_message='Document Anonymizer installed.',
    on_uninstall=_on_uninstall,
    wipe_dirs=[_JOBS_DIR],
)